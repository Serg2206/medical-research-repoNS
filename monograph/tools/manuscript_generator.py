#!/usr/bin/env python3
"""
Manuscript Generator - Scientific Medical Document Automation Tool
===================================================================

A comprehensive Python tool for generating formatted scientific medical documents
from structured markdown files, with support for multiple export formats.

Author: Medical Research Documentation System
License: MIT
Version: 1.0.0
"""

import argparse
import re
import sys
import logging
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any
from dataclasses import dataclass, field
from collections import defaultdict
import yaml

# Import libraries for document generation
try:
    from markdown2 import markdown
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from weasyprint import HTML
    from bs4 import BeautifulSoup
except ImportError as e:
    print(f"Error: Missing required library. Please install requirements: {e}")
    sys.exit(1)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


@dataclass
class DocumentSection:
    """Represents a section in the document with hierarchical structure."""
    level: int
    title: str
    content: str
    number: str = ""
    line_number: int = 0
    children: List['DocumentSection'] = field(default_factory=list)


@dataclass
class Reference:
    """Represents a bibliographic reference."""
    id: int
    text: str
    url: Optional[str] = None
    in_text_citations: List[int] = field(default_factory=list)


@dataclass
class Figure:
    """Represents a figure/image in the document."""
    id: int
    caption: str
    path: str
    reference_text: str


@dataclass
class Table:
    """Represents a table in the document."""
    id: int
    caption: str
    content: str
    reference_text: str


class ConfigManager:
    """Manages configuration settings for document generation."""
    
    DEFAULT_CONFIG = {
        'document': {
            'title': 'Scientific Manuscript',
            'author': '',
            'date': '',
            'language': 'en',
            'paper_size': 'A4',
        },
        'formatting': {
            'font_family': 'Times New Roman',
            'font_size': 12,
            'line_spacing': 1.5,
            'margins': {
                'top': 2.54,
                'bottom': 2.54,
                'left': 2.54,
                'right': 2.54
            }
        },
        'styles': {
            'heading1': {
                'font_size': 16,
                'bold': True,
                'color': '000000'
            },
            'heading2': {
                'font_size': 14,
                'bold': True,
                'color': '000000'
            },
            'heading3': {
                'font_size': 12,
                'bold': True,
                'color': '000000'
            },
            'body': {
                'font_size': 12,
                'alignment': 'justify'
            }
        },
        'numbering': {
            'sections': True,
            'figures': True,
            'tables': True,
            'format': 'decimal'  # decimal, roman, letter
        },
        'toc': {
            'enabled': True,
            'depth': 3,
            'page_break_after': True
        },
        'bibliography': {
            'enabled': True,
            'style': 'numbered',  # numbered, apa, mla
            'title': 'References'
        },
        'index': {
            'enabled': True,
            'title': 'Index'
        },
        'output': {
            'formats': ['docx', 'html', 'pdf'],
            'output_dir': './output'
        }
    }
    
    def __init__(self, config_path: Optional[Path] = None):
        """Initialize configuration manager."""
        self.config = self.DEFAULT_CONFIG.copy()
        if config_path and config_path.exists():
            self.load_config(config_path)
    
    def load_config(self, config_path: Path):
        """Load configuration from YAML file."""
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                user_config = yaml.safe_load(f)
                self._merge_config(self.config, user_config)
            logger.info(f"Configuration loaded from {config_path}")
        except Exception as e:
            logger.error(f"Error loading configuration: {e}")
            raise
    
    def _merge_config(self, base: Dict, update: Dict):
        """Recursively merge configuration dictionaries."""
        for key, value in update.items():
            if key in base and isinstance(base[key], dict) and isinstance(value, dict):
                self._merge_config(base[key], value)
            else:
                base[key] = value
    
    def get(self, key_path: str, default=None):
        """Get configuration value by dot-separated path."""
        keys = key_path.split('.')
        value = self.config
        for key in keys:
            if isinstance(value, dict) and key in value:
                value = value[key]
            else:
                return default
        return value


class MarkdownParser:
    """Parses markdown documents and extracts structured content."""
    
    def __init__(self, config: ConfigManager):
        """Initialize markdown parser."""
        self.config = config
        self.sections: List[DocumentSection] = []
        self.references: List[Reference] = []
        self.figures: List[Figure] = []
        self.tables: List[Table] = []
        self.metadata: Dict[str, str] = {}
        
    def parse(self, markdown_path: Path) -> Dict[str, Any]:
        """Parse markdown file and extract all components."""
        logger.info(f"Parsing markdown file: {markdown_path}")
        
        with open(markdown_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Extract metadata
        content = self._extract_metadata(content)
        
        # Extract references
        content, self.references = self._extract_references(content)
        
        # Parse sections
        self.sections = self._parse_sections(content)
        
        # Number sections if enabled
        if self.config.get('numbering.sections', True):
            self._number_sections(self.sections)
        
        # Extract figures and tables
        self._extract_figures_tables(content)
        
        return {
            'metadata': self.metadata,
            'sections': self.sections,
            'references': self.references,
            'figures': self.figures,
            'tables': self.tables
        }
    
    def _extract_metadata(self, content: str) -> str:
        """Extract document metadata from content."""
        # Look for metadata in bold text at the beginning
        metadata_pattern = r'\*\*(.*?):\*\*\s*(.*?)(?=\n|$)'
        matches = re.finditer(metadata_pattern, content[:1000])
        
        for match in matches:
            key = match.group(1).strip()
            value = match.group(2).strip()
            self.metadata[key] = value
        
        return content
    
    def _extract_references(self, content: str) -> Tuple[str, List[Reference]]:
        """Extract references from the document."""
        references = []
        
        # Look for references section
        ref_pattern = r'###?\s*(?:References|Источники|Bibliography)\s*\n+(.*?)(?=\n##|$)'
        match = re.search(ref_pattern, content, re.DOTALL | re.IGNORECASE)
        
        if match:
            ref_section = match.group(1)
            # Parse markdown links
            link_pattern = r'\[(.*?)\]\((.*?)\)'
            for idx, link_match in enumerate(re.finditer(link_pattern, ref_section), 1):
                references.append(Reference(
                    id=idx,
                    text=link_match.group(1),
                    url=link_match.group(2)
                ))
        
        return content, references
    
    def _parse_sections(self, content: str) -> List[DocumentSection]:
        """Parse document into hierarchical sections."""
        sections = []
        lines = content.split('\n')
        
        current_sections: Dict[int, DocumentSection] = {}
        
        for line_no, line in enumerate(lines, 1):
            # Check if line is a heading
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                title = heading_match.group(2).strip()
                
                section = DocumentSection(
                    level=level,
                    title=title,
                    content="",
                    line_number=line_no
                )
                
                # Build hierarchy
                if level == 1:
                    sections.append(section)
                    current_sections = {1: section}
                else:
                    parent_level = level - 1
                    while parent_level > 0:
                        if parent_level in current_sections:
                            current_sections[parent_level].children.append(section)
                            break
                        parent_level -= 1
                    
                    if parent_level == 0:
                        sections.append(section)
                
                current_sections[level] = section
            else:
                # Add content to the current section
                if current_sections:
                    highest_level = max(current_sections.keys())
                    if current_sections[highest_level].content:
                        current_sections[highest_level].content += '\n'
                    current_sections[highest_level].content += line
        
        return sections
    
    def _number_sections(self, sections: List[DocumentSection], prefix: str = ""):
        """Recursively number sections."""
        for idx, section in enumerate(sections, 1):
            if prefix:
                section.number = f"{prefix}.{idx}"
            else:
                section.number = str(idx)
            
            if section.children:
                self._number_sections(section.children, section.number)
    
    def _extract_figures_tables(self, content: str):
        """Extract figures and tables from content."""
        # Extract image references
        img_pattern = r'!\[(.*?)\]\((.*?)\)'
        for idx, match in enumerate(re.finditer(img_pattern, content), 1):
            self.figures.append(Figure(
                id=idx,
                caption=match.group(1),
                path=match.group(2),
                reference_text=f"Figure {idx}"
            ))
        
        # Extract tables (markdown tables)
        table_pattern = r'(\|.+\|(?:\n\|.+\|)+)'
        for idx, match in enumerate(re.finditer(table_pattern, content), 1):
            self.tables.append(Table(
                id=idx,
                caption=f"Table {idx}",
                content=match.group(1),
                reference_text=f"Table {idx}"
            ))


class DocumentGenerator:
    """Base class for document generation in various formats."""
    
    def __init__(self, config: ConfigManager, parsed_data: Dict[str, Any]):
        """Initialize document generator."""
        self.config = config
        self.parsed_data = parsed_data
        self.sections = parsed_data['sections']
        self.references = parsed_data['references']
        self.figures = parsed_data['figures']
        self.tables = parsed_data['tables']
        self.metadata = parsed_data['metadata']
    
    def generate(self, output_path: Path):
        """Generate document - to be implemented by subclasses."""
        raise NotImplementedError


class DOCXGenerator(DocumentGenerator):
    """Generates Microsoft Word (DOCX) documents."""
    
    def generate(self, output_path: Path):
        """Generate DOCX document."""
        logger.info(f"Generating DOCX document: {output_path}")
        
        doc = Document()
        
        # Set document properties
        self._setup_document_properties(doc)
        
        # Add title page
        self._add_title_page(doc)
        
        # Add table of contents
        if self.config.get('toc.enabled', True):
            self._add_toc(doc)
        
        # Add sections
        self._add_sections(doc, self.sections)
        
        # Add figures list
        if self.figures:
            self._add_figures_list(doc)
        
        # Add tables list
        if self.tables:
            self._add_tables_list(doc)
        
        # Add bibliography
        if self.config.get('bibliography.enabled', True) and self.references:
            self._add_bibliography(doc)
        
        # Save document
        doc.save(str(output_path))
        logger.info(f"DOCX document saved: {output_path}")
    
    def _setup_document_properties(self, doc: Document):
        """Set up document-wide properties."""
        # Set core properties (with 255 char limit)
        title = self.metadata.get('Цель отчета', self.config.get('document.title', 'Manuscript'))
        # Truncate title if too long (DOCX has 255 char limit for properties)
        if len(title) > 250:
            title = title[:247] + '...'
        doc.core_properties.title = title
        doc.core_properties.author = self.config.get('document.author', '')
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = self.config.get('formatting.font_family', 'Times New Roman')
        font.size = Pt(self.config.get('formatting.font_size', 12))
    
    def _add_title_page(self, doc: Document):
        """Add title page to document."""
        # Extract title from first heading or metadata
        title = None
        if self.sections and self.sections[0].level == 1:
            title = self.sections[0].title
        
        if not title:
            title = self.config.get('document.title', 'Scientific Manuscript')
        
        # Add title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        for key, value in self.metadata.items():
            p = doc.add_paragraph()
            p.add_run(f"{key}: ").bold = True
            p.add_run(value)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Page break
        doc.add_page_break()
    
    def _add_toc(self, doc: Document):
        """Add table of contents."""
        doc.add_heading('Table of Contents', level=1)
        
        # Generate TOC entries
        self._generate_toc_entries(doc, self.sections, 0)
        
        if self.config.get('toc.page_break_after', True):
            doc.add_page_break()
    
    def _generate_toc_entries(self, doc: Document, sections: List[DocumentSection], indent: int):
        """Recursively generate TOC entries."""
        max_depth = self.config.get('toc.depth', 3)
        
        for section in sections:
            if section.level <= max_depth:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5 * indent)
                
                # Add section number and title
                text = f"{section.number} {section.title}" if section.number else section.title
                p.add_run(text)
                
                # Recurse for children
                if section.children:
                    self._generate_toc_entries(doc, section.children, indent + 1)
    
    def _add_sections(self, doc: Document, sections: List[DocumentSection]):
        """Add all sections to document."""
        for section in sections:
            # Skip the title section (already added)
            if section.level == 1 and not section.number:
                continue
            
            # Add heading
            heading_text = f"{section.number} {section.title}" if section.number else section.title
            doc.add_heading(heading_text, level=min(section.level, 3))
            
            # Add content
            if section.content.strip():
                # Parse markdown formatting in content
                self._add_formatted_content(doc, section.content)
            
            # Recursively add children
            if section.children:
                self._add_sections(doc, section.children)
    
    def _add_formatted_content(self, doc: Document, content: str):
        """Add content with basic markdown formatting."""
        paragraphs = content.split('\n\n')
        
        for para_text in paragraphs:
            if not para_text.strip():
                continue
            
            p = doc.add_paragraph()
            
            # Handle bold text
            parts = re.split(r'(\*\*.*?\*\*)', para_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
    
    def _add_figures_list(self, doc: Document):
        """Add list of figures."""
        doc.add_page_break()
        doc.add_heading('List of Figures', level=1)
        
        for figure in self.figures:
            p = doc.add_paragraph()
            p.add_run(f"{figure.reference_text}: ").bold = True
            p.add_run(figure.caption)
    
    def _add_tables_list(self, doc: Document):
        """Add list of tables."""
        doc.add_page_break()
        doc.add_heading('List of Tables', level=1)
        
        for table in self.tables:
            p = doc.add_paragraph()
            p.add_run(f"{table.reference_text}: ").bold = True
            p.add_run(table.caption)
    
    def _add_bibliography(self, doc: Document):
        """Add bibliography/references section."""
        doc.add_page_break()
        bib_title = self.config.get('bibliography.title', 'References')
        doc.add_heading(bib_title, level=1)
        
        for ref in self.references:
            p = doc.add_paragraph(style='List Number')
            p.add_run(ref.text)
            if ref.url:
                p.add_run(f"\n{ref.url}").italic = True


class HTMLGenerator(DocumentGenerator):
    """Generates HTML documents."""
    
    def generate(self, output_path: Path):
        """Generate HTML document."""
        logger.info(f"Generating HTML document: {output_path}")
        
        html_content = self._build_html()
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        logger.info(f"HTML document saved: {output_path}")
    
    def _build_html(self) -> str:
        """Build complete HTML document."""
        html_parts = [
            '<!DOCTYPE html>',
            '<html lang="en">',
            '<head>',
            '    <meta charset="UTF-8">',
            '    <meta name="viewport" content="width=device-width, initial-scale=1.0">',
            f'    <title>{self.metadata.get("Цель отчета", "Manuscript")}</title>',
            '    <style>',
            self._get_css_styles(),
            '    </style>',
            '</head>',
            '<body>',
            '    <div class="container">',
        ]
        
        # Title page
        html_parts.append(self._build_title_page())
        
        # Table of contents
        if self.config.get('toc.enabled', True):
            html_parts.append(self._build_toc())
        
        # Sections
        html_parts.append(self._build_sections())
        
        # Bibliography
        if self.config.get('bibliography.enabled', True) and self.references:
            html_parts.append(self._build_bibliography())
        
        html_parts.extend([
            '    </div>',
            '</body>',
            '</html>'
        ])
        
        return '\n'.join(html_parts)
    
    def _get_css_styles(self) -> str:
        """Generate CSS styles."""
        return """
        body {
            font-family: 'Times New Roman', serif;
            font-size: 12pt;
            line-height: 1.5;
            margin: 0;
            padding: 0;
            color: #000;
        }
        .container {
            max-width: 8.5in;
            margin: 0 auto;
            padding: 1in;
        }
        h1 {
            font-size: 16pt;
            font-weight: bold;
            margin-top: 24pt;
            margin-bottom: 12pt;
            page-break-after: avoid;
        }
        h2 {
            font-size: 14pt;
            font-weight: bold;
            margin-top: 18pt;
            margin-bottom: 10pt;
            page-break-after: avoid;
        }
        h3 {
            font-size: 12pt;
            font-weight: bold;
            margin-top: 14pt;
            margin-bottom: 8pt;
            page-break-after: avoid;
        }
        p {
            text-align: justify;
            margin-bottom: 12pt;
        }
        .title-page {
            text-align: center;
            margin-bottom: 48pt;
            page-break-after: always;
        }
        .title {
            font-size: 18pt;
            font-weight: bold;
            margin-bottom: 24pt;
        }
        .metadata {
            margin-bottom: 12pt;
        }
        .toc {
            page-break-after: always;
        }
        .toc h2 {
            text-align: center;
        }
        .toc-entry {
            margin-left: 0;
            margin-bottom: 8pt;
        }
        .toc-entry.level-2 {
            margin-left: 20pt;
        }
        .toc-entry.level-3 {
            margin-left: 40pt;
        }
        .bibliography {
            page-break-before: always;
        }
        .reference {
            margin-bottom: 12pt;
            padding-left: 20pt;
            text-indent: -20pt;
        }
        @media print {
            .container {
                max-width: 100%;
            }
        }
        """
    
    def _build_title_page(self) -> str:
        """Build HTML title page."""
        title = self.sections[0].title if self.sections else "Manuscript"
        
        parts = ['<div class="title-page">']
        parts.append(f'    <h1 class="title">{title}</h1>')
        
        for key, value in self.metadata.items():
            parts.append(f'    <div class="metadata"><strong>{key}:</strong> {value}</div>')
        
        parts.append('</div>')
        
        return '\n'.join(parts)
    
    def _build_toc(self) -> str:
        """Build table of contents."""
        parts = ['<div class="toc">']
        parts.append('    <h2>Table of Contents</h2>')
        
        self._build_toc_entries(parts, self.sections)
        
        parts.append('</div>')
        
        return '\n'.join(parts)
    
    def _build_toc_entries(self, parts: List[str], sections: List[DocumentSection], level: int = 1):
        """Recursively build TOC entries."""
        max_depth = self.config.get('toc.depth', 3)
        
        for section in sections:
            if section.level <= max_depth:
                text = f"{section.number} {section.title}" if section.number else section.title
                parts.append(f'    <div class="toc-entry level-{level}">{text}</div>')
                
                if section.children:
                    self._build_toc_entries(parts, section.children, level + 1)
    
    def _build_sections(self) -> str:
        """Build all sections."""
        parts = []
        self._build_section_html(parts, self.sections)
        return '\n'.join(parts)
    
    def _build_section_html(self, parts: List[str], sections: List[DocumentSection]):
        """Recursively build section HTML."""
        for section in sections:
            # Skip title section
            if section.level == 1 and not section.number:
                continue
            
            heading_text = f"{section.number} {section.title}" if section.number else section.title
            parts.append(f'<h{min(section.level, 3)}>{heading_text}</h{min(section.level, 3)}>')
            
            if section.content.strip():
                # Convert markdown to HTML
                content_html = markdown(section.content)
                parts.append(content_html)
            
            if section.children:
                self._build_section_html(parts, section.children)
    
    def _build_bibliography(self) -> str:
        """Build bibliography section."""
        parts = ['<div class="bibliography">']
        bib_title = self.config.get('bibliography.title', 'References')
        parts.append(f'    <h2>{bib_title}</h2>')
        
        for ref in self.references:
            parts.append(f'    <div class="reference">')
            parts.append(f'        <strong>[{ref.id}]</strong> {ref.text}')
            if ref.url:
                parts.append(f'        <br><em><a href="{ref.url}">{ref.url}</a></em>')
            parts.append('    </div>')
        
        parts.append('</div>')
        
        return '\n'.join(parts)


class PDFGenerator(HTMLGenerator):
    """Generates PDF documents via HTML conversion."""
    
    def generate(self, output_path: Path):
        """Generate PDF document."""
        logger.info(f"Generating PDF document: {output_path}")
        
        html_content = self._build_html()
        
        # Convert HTML to PDF using WeasyPrint
        HTML(string=html_content).write_pdf(str(output_path))
        
        logger.info(f"PDF document saved: {output_path}")


class ManuscriptGenerator:
    """Main class orchestrating the document generation process."""
    
    def __init__(self, config_path: Optional[Path] = None):
        """Initialize manuscript generator."""
        self.config = ConfigManager(config_path)
        self.parser = MarkdownParser(self.config)
    
    def generate(self, input_path: Path, output_dir: Optional[Path] = None):
        """Generate documents from markdown input."""
        # Parse markdown
        parsed_data = self.parser.parse(input_path)
        
        # Determine output directory
        if output_dir is None:
            output_dir = Path(self.config.get('output.output_dir', './output'))
        
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Get base filename
        base_name = input_path.stem
        
        # Generate documents in requested formats
        formats = self.config.get('output.formats', ['docx', 'html', 'pdf'])
        
        generators = {
            'docx': DOCXGenerator,
            'html': HTMLGenerator,
            'pdf': PDFGenerator
        }
        
        generated_files = []
        
        for fmt in formats:
            if fmt in generators:
                output_path = output_dir / f"{base_name}.{fmt}"
                generator = generators[fmt](self.config, parsed_data)
                
                try:
                    generator.generate(output_path)
                    generated_files.append(output_path)
                except Exception as e:
                    logger.error(f"Error generating {fmt} format: {e}")
                    raise
        
        logger.info(f"Document generation complete. Files: {generated_files}")
        return generated_files


def main():
    """Main CLI entry point."""
    parser = argparse.ArgumentParser(
        description='Manuscript Generator - Automate scientific document generation',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Generate with default configuration
  %(prog)s input.md
  
  # Generate with custom configuration
  %(prog)s input.md -c config.yaml
  
  # Generate specific formats only
  %(prog)s input.md -f docx pdf
  
  # Specify output directory
  %(prog)s input.md -o ./output
        """
    )
    
    parser.add_argument(
        'input',
        type=Path,
        help='Input markdown file path'
    )
    
    parser.add_argument(
        '-c', '--config',
        type=Path,
        help='Configuration file path (YAML)'
    )
    
    parser.add_argument(
        '-o', '--output',
        type=Path,
        help='Output directory'
    )
    
    parser.add_argument(
        '-f', '--formats',
        nargs='+',
        choices=['docx', 'html', 'pdf'],
        help='Output formats (default: all)'
    )
    
    parser.add_argument(
        '-v', '--verbose',
        action='store_true',
        help='Enable verbose logging'
    )
    
    args = parser.parse_args()
    
    # Set logging level
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    # Validate input file
    if not args.input.exists():
        logger.error(f"Input file not found: {args.input}")
        sys.exit(1)
    
    # Create generator
    try:
        generator = ManuscriptGenerator(args.config)
        
        # Override formats if specified
        if args.formats:
            generator.config.config['output']['formats'] = args.formats
        
        # Generate documents
        generated_files = generator.generate(args.input, args.output)
        
        print("\n✓ Document generation successful!")
        print("\nGenerated files:")
        for file in generated_files:
            print(f"  - {file}")
        
    except Exception as e:
        logger.error(f"Document generation failed: {e}")
        if args.verbose:
            raise
        sys.exit(1)


if __name__ == '__main__':
    main()
